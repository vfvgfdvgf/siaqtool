import tempfile
import unittest
import zipfile
import json
from pathlib import Path

from PIL import Image
from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader, PdfWriter

from converter.processors import SUPPORTED_TOOLS, process
from converter.security import ALLOWED_EXTENSIONS

class ProcessorTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory(); self.root = Path(self.temp.name)

    def tearDown(self): self.temp.cleanup()

    def pdf(self, name: str, pages: int) -> Path:
        path = self.root / name; writer = PdfWriter()
        for _ in range(pages): writer.add_blank_page(width=300, height=400)
        with path.open("wb") as target: writer.write(target)
        return path

    def test_merge_pdf(self):
        result = process("merge-pdf", [self.pdf("one.pdf", 1), self.pdf("two.pdf", 2)], {}, self.root)
        self.assertEqual(len(PdfReader(result.path).pages), 3)

    def test_split_pdf(self):
        result = process("split-pdf", [self.pdf("source.pdf", 3)], {}, self.root)
        with zipfile.ZipFile(result.path) as bundle: self.assertEqual(len(bundle.namelist()), 3)

    def test_rotate_pdf(self):
        result = process("rotate-pdf", [self.pdf("source.pdf", 1)], {"angle": 90}, self.root)
        page = PdfReader(result.path).pages[0]
        self.assertEqual(page.rotation, 90)

    def test_images_to_pdf(self):
        image_path = self.root / "image.jpg"; Image.new("RGB", (120, 80), "white").save(image_path)
        result = process("jpg-to-pdf", [image_path], {}, self.root)
        self.assertEqual(len(PdfReader(result.path).pages), 1)

    def test_new_tool_catalog_is_wired(self):
        self.assertEqual(len(SUPPORTED_TOOLS), 145)
        self.assertEqual(SUPPORTED_TOOLS, frozenset(ALLOWED_EXTENSIONS))

    def test_delete_pdf_page(self):
        result = process("delete-pdf-pages", [self.pdf("source.pdf", 3)], {"pages": "2"}, self.root)
        self.assertEqual(len(PdfReader(result.path).pages), 2)

    def test_excel_to_word(self):
        source = self.root / "source.xlsx"
        workbook = Workbook(); workbook.active.append(["الاسم", "القيمة"]); workbook.active.append(["سياق", 100]); workbook.save(source)
        result = process("excel-to-word", [source], {}, self.root)
        document = Document(result.path)
        self.assertEqual(document.tables[0].cell(1, 0).text, "سياق")

    def test_word_to_excel(self):
        source = self.root / "source.docx"
        document = Document(); document.add_paragraph("نص تجريبي"); document.save(source)
        result = process("word-to-excel", [source], {}, self.root)
        workbook = load_workbook(result.path)
        self.assertEqual(workbook["النص"]["A1"].value, "نص تجريبي")

    def test_png_to_jpg(self):
        source = self.root / "source.png"; Image.new("RGBA", (80, 60), (255, 0, 0, 120)).save(source)
        result = process("png-to-jpg", [source], {}, self.root)
        with Image.open(result.path) as converted:
            self.assertEqual(converted.format, "JPEG")

    def test_pdf_info(self):
        result = process("pdf-info", [self.pdf("source.pdf", 2)], {}, self.root)
        payload = json.loads(result.path.read_text(encoding="utf-8"))
        self.assertEqual(payload["pages"], 2)

    def test_csv_to_json(self):
        source = self.root / "source.csv"
        source.write_text("name,value\nSiaq,146\n", encoding="utf-8")
        result = process("csv-to-json", [source], {}, self.root)
        payload = json.loads(result.path.read_text(encoding="utf-8"))
        self.assertEqual(payload["rows"][0]["name"], "Siaq")

    def test_image_palette(self):
        source = self.root / "palette.png"
        Image.new("RGB", (20, 20), "#d9232e").save(source)
        result = process("image-palette", [source], {}, self.root)
        payload = json.loads(result.path.read_text(encoding="utf-8"))
        self.assertEqual(payload["colors"][0]["hex"], "#d9232e")

if __name__ == "__main__": unittest.main()
