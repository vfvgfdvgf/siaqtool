from __future__ import annotations

import csv
import json
import zipfile
from pathlib import Path

import pymupdf as fitz
from docx import Document
from openpyxl import Workbook, load_workbook

from .core import ProcessResult, ProcessingError, file_result, libreoffice_convert


def _read_delimited(path: Path, delimiter: str | None = None) -> tuple[list[str], list[dict[str, str]]]:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    if delimiter is None:
        try:
            delimiter = csv.Sniffer().sniff(text[:8192], delimiters=",;\t|").delimiter
        except csv.Error:
            delimiter = ","
    reader = csv.DictReader(text.splitlines(), delimiter=delimiter)
    raw_headers = [header for header in (reader.fieldnames or []) if header and header.strip()]
    headers = [header.strip() for header in raw_headers]
    if not headers:
        raise ProcessingError("لم نجد صف عناوين صالحًا في الملف.")
    rows = [
        {header.strip(): str(row.get(header, "") or "") for header in raw_headers}
        for row in reader
    ]
    return headers, rows


def _write_delimited(path: Path, headers: list[str], rows: list[dict], delimiter: str) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as target:
        writer = csv.DictWriter(target, fieldnames=headers, delimiter=delimiter, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def csv_to_json(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    headers, rows = _read_delimited(paths[0])
    output = workdir / "table.json"
    output.write_text(json.dumps({"columns": headers, "rows": rows}, ensure_ascii=False, indent=2), encoding="utf-8")
    return file_result(output, "siaq-table.json", "application/json")


def _json_rows(path: Path) -> tuple[list[str], list[dict]]:
    try:
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
    except (json.JSONDecodeError, UnicodeDecodeError) as exc:
        raise ProcessingError("ملف JSON غير صالح.") from exc
    if isinstance(payload, dict) and isinstance(payload.get("rows"), list):
        payload = payload["rows"]
    elif isinstance(payload, dict):
        payload = [payload]
    if not isinstance(payload, list) or not payload:
        raise ProcessingError("يجب أن يحتوي JSON على قائمة سجلات أو كائن واحد.")
    rows = [item if isinstance(item, dict) else {"value": item} for item in payload]
    headers = list(dict.fromkeys(str(key) for row in rows for key in row))
    return headers, [{header: row.get(header, "") for header in headers} for row in rows]


def json_to_csv(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    headers, rows = _json_rows(paths[0])
    output = workdir / "table.csv"
    _write_delimited(output, headers, rows, ",")
    return file_result(output, "siaq-table.csv", "text/csv; charset=utf-8")


def json_to_excel(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    headers, rows = _json_rows(paths[0])
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "البيانات"
    sheet.append(headers)
    for row in rows:
        sheet.append([row.get(header, "") if not isinstance(row.get(header), (dict, list)) else json.dumps(row.get(header), ensure_ascii=False) for header in headers])
    output = workdir / "table.xlsx"
    workbook.save(output)
    return file_result(output, "siaq-table.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def excel_to_json(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    source = paths[0] if paths[0].suffix.lower() == ".xlsx" else libreoffice_convert(paths[0], workdir, "xlsx")
    workbook = load_workbook(source, read_only=True, data_only=True)
    sheets = []
    for sheet in workbook.worksheets:
        values = list(sheet.iter_rows(values_only=True))
        if not values:
            continue
        headers = [str(value or f"column_{index}") for index, value in enumerate(values[0], start=1)]
        rows = [dict(zip(headers, row)) for row in values[1:]]
        sheets.append({"name": sheet.title, "columns": headers, "rows": rows})
    workbook.close()
    output = workdir / "workbook.json"
    output.write_text(json.dumps({"sheets": sheets}, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return file_result(output, "siaq-workbook.json", "application/json")


def convert_delimiter(paths: list[Path], workdir: Path, _options: dict, source: str | None, target: str) -> ProcessResult:
    headers, rows = _read_delimited(paths[0], source)
    extension = "tsv" if target == "\t" else "csv"
    output = workdir / f"table.{extension}"
    _write_delimited(output, headers, rows, target)
    content_type = "text/tab-separated-values" if extension == "tsv" else "text/csv"
    return file_result(output, f"siaq-table.{extension}", f"{content_type}; charset=utf-8")


def merge_csv(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    if len(paths) < 2:
        raise ProcessingError("دمج CSV يحتاج ملفين على الأقل.")
    datasets = [_read_delimited(path) for path in paths]
    headers = list(dict.fromkeys(header for dataset_headers, _ in datasets for header in dataset_headers))
    rows = [row for _, dataset_rows in datasets for row in dataset_rows]
    output = workdir / "merged.csv"
    _write_delimited(output, headers, rows, ",")
    return file_result(output, "siaq-merged.csv", "text/csv; charset=utf-8")


def split_csv(paths: list[Path], workdir: Path, options: dict) -> ProcessResult:
    headers, rows = _read_delimited(paths[0])
    chunk_size = max(100, min(int(options.get("rows", 500)), 5000))
    archive = workdir / "csv-parts.zip"
    with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as bundle:
        chunks = [rows[start:start + chunk_size] for start in range(0, len(rows), chunk_size)] or [[]]
        for index, chunk in enumerate(chunks, start=1):
            part = workdir / f"part-{index}.csv"
            _write_delimited(part, headers, chunk, ",")
            bundle.write(part, part.name)
    return file_result(archive, "siaq-csv-parts.zip", "application/zip")


def deduplicate_csv(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    headers, rows = _read_delimited(paths[0])
    seen: set[tuple[str, ...]] = set()
    unique = []
    for row in rows:
        key = tuple(row.get(header, "").strip() for header in headers)
        if key not in seen:
            seen.add(key)
            unique.append(row)
    output = workdir / "unique.csv"
    _write_delimited(output, headers, unique, ",")
    return file_result(output, "siaq-unique.csv", "text/csv; charset=utf-8")


def csv_summary(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    headers, rows = _read_delimited(paths[0])
    empty = sum(1 for row in rows for header in headers if not row.get(header, "").strip())
    payload = {"columns": len(headers), "rows": len(rows), "headers": headers, "empty_cells": empty, "filled_cells": len(headers) * len(rows) - empty}
    output = workdir / "csv-summary.json"
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return file_result(output, "siaq-csv-summary.json", "application/json")


def text_to_pdf(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    output = libreoffice_convert(paths[0], workdir, "pdf")
    return file_result(output, "siaq-text.pdf", "application/pdf")


def markdown_to_docx(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    document = Document()
    for raw in paths[0].read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.strip()
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line:
            document.add_paragraph(line)
    output = workdir / "markdown.docx"
    document.save(output)
    return file_result(output, "siaq-markdown.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def office_to_html(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    output = libreoffice_convert(paths[0], workdir, "html")
    return file_result(output, "siaq-document.html", "text/html; charset=utf-8")


def powerpoint_to_images(paths: list[Path], workdir: Path, _options: dict) -> ProcessResult:
    pdf = libreoffice_convert(paths[0], workdir, "pdf")
    document = fitz.open(pdf)
    archive = workdir / "slides.zip"
    try:
        with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as bundle:
            for index, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = workdir / f"slide-{index}.png"
                pixmap.save(image)
                bundle.write(image, image.name)
    finally:
        document.close()
    return file_result(archive, "siaq-slides.zip", "application/zip")


DATA_PROCESSORS = {
    "csv-to-json": csv_to_json,
    "json-to-csv": json_to_csv,
    "json-to-excel": json_to_excel,
    "excel-to-json": excel_to_json,
    "csv-to-tsv": lambda p, w, o: convert_delimiter(p, w, o, None, "\t"),
    "tsv-to-csv": lambda p, w, o: convert_delimiter(p, w, o, "\t", ","),
    "merge-csv": merge_csv,
    "split-csv": split_csv,
    "deduplicate-csv": deduplicate_csv,
    "csv-summary": csv_summary,
    "text-to-pdf": text_to_pdf,
    "markdown-to-docx": markdown_to_docx,
    "docx-to-html": office_to_html,
    "xlsx-to-html": office_to_html,
    "powerpoint-to-images": powerpoint_to_images,
}
